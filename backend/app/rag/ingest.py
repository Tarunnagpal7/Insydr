import os
import csv
import asyncio
from io import StringIO
from pathlib import Path
from typing import List, Dict, Any
from pypdf import PdfReader
from app.rag.chunker import chunk_text
from app.rag.embeddings import EmbeddingService

# Supported file extensions and their MIME types
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.csv', '.md'}
SUPPORTED_MIMES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/plain',
    'text/csv',
    'text/markdown',
    'application/octet-stream',  # Fallback for .md files
}


class IngestionPipeline:
    def __init__(self):
        self.embedding_service = EmbeddingService()

    async def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Asynchronously processes a document by running the synchronous logic in a thread.
        """
        return await asyncio.to_thread(self._process_sync, file_path)

    async def process_text(self, text: str, title: str = "Untitled") -> Dict[str, Any]:
        """
        Process raw text content (for manual text entry).
        """
        return await asyncio.to_thread(self._process_text_sync, text, title)

    def _process_text_sync(self, text: str, title: str) -> Dict[str, Any]:
        """Chunk and embed raw text."""
        text_chunks = chunk_text(text)
        chunks_data = self._embed_chunks(text_chunks)
        return {"title": title, "chunks": chunks_data}

    def _process_sync(self, file_path: str) -> Dict[str, Any]:
        print(f"[DEBUG] Starting ingestion for {file_path}")
        """
        Extracts text, chunks it, and generates embeddings.
        Returns a dict with 'title', 'chunks' (content, embedding, index, tokens).
        """
        # 1. Extract Text
        text = self._extract_text(file_path)
        title = Path(file_path).stem

        # 2. Chunk Text
        text_chunks = chunk_text(text)

        # 3. Embed Chunks
        chunks_data = self._embed_chunks(text_chunks)

        return {
            "title": title,
            "chunks": chunks_data
        }

    def _embed_chunks(self, text_chunks: List[str]) -> List[Dict[str, Any]]:
        """Batch-embed text chunks and return structured chunk data."""
        batch_size = 32
        embeddings = []

        for i in range(0, len(text_chunks), batch_size):
            batch = text_chunks[i : i + batch_size]
            try:
                batch_embeddings = self.embedding_service.embed_documents(batch)
                embeddings.extend(batch_embeddings)
            except Exception as e:
                print(f"Error embedding batch {i}: {e}")
                raise e

        chunks_data = []
        for i, chunk_content in enumerate(text_chunks):
            # approximate token count
            token_count = len(chunk_content.split())

            # Ensure embedding exists
            embedding_vector = embeddings[i] if i < len(embeddings) else [0.0] * 384

            chunks_data.append({
                "content": chunk_content,
                "embedding": embedding_vector,
                "chunk_index": i,
                "token_count": token_count
            })
        return chunks_data

    def _extract_text(self, file_path: str) -> str:
        """Route extraction based on file extension."""
        ext = Path(file_path).suffix.lower()
        
        extractors = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.txt': self._extract_txt,
            '.md': self._extract_md,
            '.csv': self._extract_csv,
        }

        extractor = extractors.get(ext)
        if extractor is None:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
        
        return extractor(file_path)

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF preserving page structure."""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
            raise e
        return text.strip()

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX preserving headings and paragraph structure."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        text_parts = []
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                # Preserve heading structure with markdown-style markers
                if para.style and para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '').replace('Heading', '1')
                    try:
                        level_num = int(level)
                    except ValueError:
                        level_num = 1
                    text_parts.append(f"{'#' * level_num} {para.text.strip()}")
                else:
                    text_parts.append(para.text.strip())
            
            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            raise e
        
        return "\n\n".join(text_parts)

    def _extract_txt(self, file_path: str) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read().strip()
        except Exception as e:
            print(f"Error reading TXT: {e}")
            raise e

    def _extract_md(self, file_path: str) -> str:
        """Extract text from Markdown files, preserving structure."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read().strip()
        except Exception as e:
            print(f"Error reading Markdown: {e}")
            raise e

    def _extract_csv(self, file_path: str) -> str:
        """
        Extract from CSV, converting each row into structured text blocks.
        Format: 'Column1: Value1, Column2: Value2, ...'
        Great for FAQ-style data or structured datasets.
        """
        text_parts = []
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.DictReader(f)
                if reader.fieldnames is None:
                    # Fallback: read as plain text  
                    f.seek(0)
                    return f.read().strip()
                
                for row in reader:
                    parts = []
                    for key, value in row.items():
                        if key and value and value.strip():
                            parts.append(f"{key}: {value.strip()}")
                    if parts:
                        text_parts.append(". ".join(parts))
        except Exception as e:
            print(f"Error reading CSV: {e}")
            raise e
        
        return "\n\n".join(text_parts)
